from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TITLE_MAX_LENGTH = 40
BODY_FONT_NAME = "仿宋_GB2312"
TITLE_FONT_NAME = "方正小标宋简体"
LEVEL1_FONT_NAME = "黑体"
LEVEL2_FONT_NAME = "楷体_GB2312"


def export_official_docx(content: str, target_path: str | Path) -> Path:
    document = Document()
    _configure_document(document)
    _remove_default_paragraph(document)

    normalized_lines = _normalize_lines(content)
    if not normalized_lines:
        normalized_lines = [""]

    title_text, body_lines = _split_title_and_body(normalized_lines)
    if title_text:
        _add_title_paragraph(document, title_text)

    for line in body_lines:
        if not line.strip():
            document.add_paragraph("")
            continue
        _add_body_paragraph(document, line)

    resolved_target = Path(target_path).expanduser().resolve()
    resolved_target.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(resolved_target))
    return resolved_target


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal_style = document.styles["Normal"]
    normal_style.font.name = BODY_FONT_NAME
    normal_style.font.size = Pt(16)
    normal_style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT_NAME)


def _remove_default_paragraph(document: Document) -> None:
    if len(document.paragraphs) != 1:
        return
    paragraph = document.paragraphs[0]
    if paragraph.text.strip():
        return
    paragraph._element.getparent().remove(paragraph._element)


def _normalize_lines(content: str) -> list[str]:
    lines: list[str] = []
    previous_blank = False
    for raw_line in str(content or "").replace("\r\n", "\n").split("\n"):
        line = str(raw_line).rstrip()
        if not line.strip():
            if previous_blank:
                continue
            lines.append("")
            previous_blank = True
            continue
        lines.append(line.strip())
        previous_blank = False
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    return lines


def _split_title_and_body(lines: list[str]) -> tuple[str, list[str]]:
    if not lines:
        return "", []

    first_line = _strip_markdown_heading(lines[0])
    if _looks_like_title(first_line):
        return first_line, lines[1:]
    return "", lines


def _looks_like_title(text: str) -> bool:
    normalized = str(text or "").strip()
    if not normalized or len(normalized) > TITLE_MAX_LENGTH:
        return False
    if normalized.startswith(("一、", "（一）")):
        return False
    if re.match(r"^[0-9]+[.．、]", normalized):
        return False
    if normalized.endswith(("。", "！", "？", "；", "：", ":", ";", ".", "!", "?")):
        return False
    return True


def _strip_markdown_heading(text: str) -> str:
    normalized = str(text or "").strip()
    if normalized.startswith("#"):
        return normalized.lstrip("#").strip()
    return normalized


def _add_title_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(18)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(28)
    run = paragraph.add_run(text)
    _set_run_font(run, TITLE_FONT_NAME, 22)


def _add_body_paragraph(document: Document, text: str) -> None:
    normalized = _strip_markdown_heading(text)
    paragraph = document.add_paragraph()
    paragraph.alignment = _resolve_alignment(normalized)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(28)
    paragraph.paragraph_format.first_line_indent = Pt(32)

    run = paragraph.add_run(normalized)
    font_name = BODY_FONT_NAME
    if _is_level1_heading(normalized):
        font_name = LEVEL1_FONT_NAME
    elif _is_level2_heading(normalized):
        font_name = LEVEL2_FONT_NAME
    _set_run_font(run, font_name, 16)

    if paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        paragraph.paragraph_format.first_line_indent = Pt(0)


def _resolve_alignment(text: str) -> WD_ALIGN_PARAGRAPH:
    if _looks_like_date_line(text):
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.JUSTIFY


def _looks_like_date_line(text: str) -> bool:
    normalized = str(text or "").strip()
    return bool(re.match(r"^[0-9〇零一二三四五六七八九十]{2,4}年[0-9一二三四五六七八九十]{1,2}月[0-9一二三四五六七八九十]{1,3}日$", normalized))


def _is_level1_heading(text: str) -> bool:
    return bool(re.match(r"^[一二三四五六七八九十]+、", str(text or "").strip()))


def _is_level2_heading(text: str) -> bool:
    return bool(re.match(r"^（[一二三四五六七八九十]+）", str(text or "").strip()))


def _set_run_font(run: object, font_name: str, font_size: int) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
